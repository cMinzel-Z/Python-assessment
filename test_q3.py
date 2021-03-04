# -*- coding: utf-8 -*-
"""
Created on Tue Nov 12 11:16:12 2019

@author: Minzel
"""

import docx
import os, sys

def change_style(filepath, style='IEEE'):
    from docx import Document
    document = Document(filepath)
   
    if style == 'IEEE':
        # find index of references paragraph
        def findRefer():
            n = 0
            for para in document.paragraphs: 
                if para.text != 'References':
                    n = n + 1
                else:
                    return n
                
        refer_para = findRefer()
        
        # extract in-text citation into a list
        def getCitation(test_p):
            index = 0
            startIndex = []
            while index < len(test_p[:refer_para]):
                index = test_p.find('(', index)
                if index == -1:
                    break
                startIndex.append(index)
                index += 1
            
            index2 = 0
            endIndex = []
            while index2 < len(test_p[:refer_para]):
                index2 = test_p.find(')', index2)
                if index2 == -1:
                    break
                endIndex.append(index2)
                index2 += 1
        
            n = 0
            test_p_l = []
            for n in range(0,len(startIndex)):
                if n < len(startIndex):
                    a = startIndex[n]
                    b = endIndex[n]
                    test_p_l.append(test_p[a:(b+1)])
                    n = n + 1
            return test_p_l
    
        # get paragraphs which include in-text citation
        def getParagraph():
            i = 0
            para_l = []
            for i in range(len(document.paragraphs[:refer_para])):
                if '(' and ')' in document.paragraphs[i].text:
                    para_l.append(document.paragraphs[i].text)
                else:
                    para_l.append('no citation')
                i = i + 1
            return para_l
        para_l = getParagraph()
        
        # get citation list with dirty data
        def getCitaL():
            i = 0
            cita_l = []
            for i in range(len(para_l)):
                cita_l.append(getCitation(para_l[i]))
                i = i + 1
            return cita_l
        cita_l = getCitaL()
        
        # find out the index of paragraph which has citation
        def findCitaPara():
            para_l_c = []
            i = 0
            for i in range(len(cita_l)):
                if '(' in str(cita_l[i]):
                    para_l_c.append(i)
                    i = i + 1
            return para_l_c
        
        # a list includes index of 'in-text' citation paragraphs  
        para_l_c = findCitaPara()
        
        
        # extract references list citation into a list
        def getCitation_r(test_p):
            index = 0
            startIndex = []
            while index < len(test_p[refer_para:]):
                index = test_p.find('(', index)
                if index == -1:
                    break
                startIndex.append(index)
                index += 1
            
            index2 = 0
            endIndex = []
            while index2 < len(test_p[refer_para:]):
                index2 = test_p.find(')', index2)
                if index2 == -1:
                    break
                endIndex.append(index2)
                index2 += 1
        
            n = 0
            test_p_l_r = []
            for n in range(0,len(startIndex)):
                if n < len(startIndex):
                    a = startIndex[n]
                    b = endIndex[n]
                    test_p_l_r.append(test_p[a:(b+1)])
                    n = n + 1
            return test_p_l_r
    
        # get paragraphs which include references list citation
        def getParagraph_r():
            i = 0
            para_l_r = []
            for i in range(len(document.paragraphs[refer_para:])):
                if '(' and ')' in document.paragraphs[-i].text:
                    para_l_r.append(document.paragraphs[-i].text)
                else:
                    para_l_r.append('no citation')
                i = i + 1
            return para_l_r
        para_l_r = getParagraph_r()
        
        # get citation list with dirty data
        def getCitaL_r():
            i = 0
            cita_l_r = []
            for i in range(len(para_l_r)):
                cita_l_r.append(getCitation_r(para_l_r[i]))
                i = i + 1
            return cita_l_r
        cita_l_r = getCitaL_r()
        
        # find out the index of paragraph which has citation
        def findCitaPara_r():
            para_l_c_r = []
            i = 0
            for i in range(len(cita_l_r)):
                if '(' in str(cita_l_r[i]):
                    para_l_c_r.append(i)
                    i = i + 1
            return para_l_c_r
        
        # a list includes index of 'references list' citation paragraphs  
        para_l_c_r = findCitaPara_r()
        para_l_c_r = [ -x for x in para_l_c_r]
        
        # add citation number to each reference
        each_refer = 0
        cita_num = len(para_l_c_r)
        citation_year = []
        citation_name = []
        for each_refer in range(len(para_l_c_r)):
            # citation number
            document.paragraphs[para_l_c_r[each_refer]].text = str(cita_num) + '. '+ document.paragraphs[para_l_c_r[each_refer]].text
            # move (date) to the end
            brackets_l = document.paragraphs[para_l_c_r[each_refer]].text.find('(')
            brackets_r = document.paragraphs[para_l_c_r[each_refer]].text.find(')')
            move_date = document.paragraphs[para_l_c_r[each_refer]].text[brackets_l:(brackets_r+3)]
            document.paragraphs[para_l_c_r[each_refer]].text = document.paragraphs[para_l_c_r[each_refer]].text.replace(move_date, '')
            document.paragraphs[para_l_c_r[each_refer]].text = document.paragraphs[para_l_c_r[each_refer]].text + ' ' + move_date.replace('(', '').replace(')', '')
            # find citation year
            citation_year.append(document.paragraphs[para_l_c_r[each_refer]].text[-6:-2])
            # find citation author name
            name_Index_l = document.paragraphs[para_l_c_r[each_refer]].text.find(' ')
            name_Index_r = document.paragraphs[para_l_c_r[each_refer]].text.find(',')
            citation_name.append(document.paragraphs[para_l_c_r[each_refer]].text[(name_Index_l+1):name_Index_r])
            each_refer = each_refer + 1
            cita_num = cita_num - 1
        # the loop is reversed, we need to get a forward sorted list
        citation_year.reverse()
        citation_name.reverse()
        
        # find all subset and return index
        def find_all(source,dest):
            length1,length2 = len(source),len(dest)
            dest_list = []
            temp_list = []
            if length1 < length2:
                return -1
            i = 0
            while i <= length1-length2:
                if source[i] == dest[0]:
                    dest_list.append(i)
                i += 1
            if dest_list == []:
                return -1
            for x in dest_list:
                if source[x:x+length2] != dest:
                    temp_list.append(x)
            for x in temp_list:
                dest_list.remove(x)
            return dest_list    
        
        # get capital letters index
        def getIndices(s):
            return [i for i, c in enumerate(s) if c.isupper()]
        
        # change in-text citation
        def changeInTextC():
            citation_content = []
            for i in range(len(para_l_c)):
                para_text = document.paragraphs[para_l_c[i]].text
                name_IN_Index = find_all(para_text,'(')
                year_IN_Index = find_all(para_text, ')')
                for x, y in zip(name_IN_Index, year_IN_Index):
                    if para_text[(x+1):y].isdigit():
                        if x >= 23:
                            citation_content.append(para_text[(x-23):y])
                        else:
                            citation_content.append(para_text[(x-12):y])
                    else:
                        citation_content.append(para_text[x:y])
            return citation_content
                    
        citation_content = changeInTextC()
        
        # add missing right brackets
        def addRightBrackets():
            for i in range(len(citation_content)):
                citation_content[i] = citation_content[i] + ')'
            return citation_content
        
        citation_content = addRightBrackets()
        
        # keep author name and date
        def rmDirtyCita():
            author_name_Index = []
            for i in citation_content:
                if i[0] != '(':
                    author_name_Index.append(getIndices(i))
                else:
                    author_name_Index.append('')
            return author_name_Index
        author_name_Index = rmDirtyCita()
        
        def rmDirtyCita2():
            i = 0
            for i in range(len(author_name_Index)):
                if type(author_name_Index[i]) == list:
                    if len(author_name_Index[i]) == 2:
                        author_name_Index[i].pop()
                i = i + 1
            return author_name_Index
        author_name_Index = rmDirtyCita2()
        
        def rmDirtyCita3():
            i = 0
            for i in range(len(author_name_Index)):
                if type(author_name_Index[i]) == list:
                    author_name_Index[i] = int(str(author_name_Index[i]).replace('[','').replace(']',''))
                i = i + 1
            return author_name_Index
        author_name_Index = rmDirtyCita3()
        author_name_Index_In = list(range(len(author_name_Index)))
        dict_author_dir = dict(zip(author_name_Index_In,author_name_Index))
        
        def needRmCita():
            needRmCita = []
            for key in dict_author_dir:
                if type(dict_author_dir[key]) == int:
                    needRmCita.append(key)
            return needRmCita
        needRmCita = needRmCita()
        
        def needRmCitaDict():
            needRmCitaDict = []
            for i in needRmCita:
                needRmCitaDict.append(author_name_Index[i])
            return needRmCitaDict
        needRmCitaDict = needRmCitaDict()
        
        rm_cita_dict = dict(zip(needRmCita, needRmCitaDict))
        
        def rmLeftNotNeed():
            for key in rm_cita_dict:
                citation_content[key] = citation_content[key][rm_cita_dict[key]:]
                
        rmLeftNotNeed()
    
        def citaFilter():
            filt_cita_content = []
            for i in citation_content:
                if ', ' in i:
                    if i[-5:-1].isdigit():
                        if ';' in i:
                            filt_cita_content.append(i.lstrip('(').rstrip(')').split('; '))
                        else:
                            filt_cita_content.append(i)
                elif i[-5:-1].isdigit():
                    filt_cita_content.append(i)
            return filt_cita_content
        filt_cita_content = citaFilter()
                         
        def flatten(A):
            rt = []
            for i in A:
                if isinstance(i,list): rt.extend(flatten(i))
                else: rt.append(i)
            return rt
        
        filt_cita_content = flatten(filt_cita_content)
        # remove duplicates
        # will be used for replacing in-text citation
        mylist = list(dict.fromkeys(filt_cita_content))
        
        para_r_l = para_l_c_r
        para_r_l.reverse()
        
        
        def rmNoCita(x):
            if 'no citation' in x: 
                x.remove('no citation')
            return x
    
        para_l_r = rmNoCita(para_l_r)
        para_l_r = rmNoCita(para_l_r)
        
        # transfering style for in-text format
        def referCitaTrans():
            n = 1
            mydict = {}
            for i in para_l_r:
                citation_need_trans_r_index = i.find(')')
                if i[:(citation_need_trans_r_index+1)].count(',') == 1:
                    comma_Index = i.find(',')
                    brackts_l = i.find('(')
                    brackts_r = i.find(')')
                    part_1 = i[:comma_Index]
                    part_2 = i[(brackts_l+1):brackts_r]
                    cita_style_1 = part_1 + ', ' + part_2
                    cita_style_2 = '(' + part_1 + ', ' + part_2 + ')'
                    cita_style_3 = part_1 + ' ' + '(' + part_2 + ')'
                    mydict[cita_style_1] = n
                    mydict[cita_style_2] = n
                    mydict[cita_style_3] = n
                elif i[:(citation_need_trans_r_index+1)].count(',') == 3:
                    comma_l = find_all(i, ',')
                    comma_1 = comma_l[0]
                    comma_2 = comma_l[2]
                    brackts_l = i.find('(')
                    brackts_r = i.find(')')
                    and_Index = i.find(' and ')
                    second_name_in = and_Index + 5
                    part_1 = i[:comma_1]
                    part_2 = i[second_name_in:(comma_2+1)]
                    part_3 = i[(brackts_l+1):brackts_r]
                    cita_style_1 = part_1 + ' and ' + part_2 + ' ' + part_3
                    cita_style_2 = '(' + part_1 + ' and ' + part_2 + ' ' + part_3 + ')'
                    cita_style_3 = part_1 + ' and ' + part_2 + ' ' + '(' + part_3 + ')'
                    mydict[cita_style_1] = n
                    mydict[cita_style_2] = n
                    mydict[cita_style_3] = n
                elif i[:(citation_need_trans_r_index+1)].count(',') >= 5:
                    comma_Index = i.find(',')
                    brackts_l = i.find('(')
                    brackts_r = i.find(')')
                    part_1 = i[:comma_Index]
                    part_2 = i[(brackts_l+1):brackts_r]
                    cita_style_1 = part_1 + ' et al., ' + part_2
                    cita_style_2 = '(' + part_1 + ' et al., ' + part_2 + ')'
                    cita_style_3 = part_1 + ' et al. ' + '(' + part_2 + ')'
                    mydict[cita_style_1] = n
                    mydict[cita_style_2] = n
                    mydict[cita_style_3] = n
                n = n + 1
            return mydict
        mydict = referCitaTrans()
        
        # remove interference information
        def rmEG(mylist):
            i = 0
            for i in range(len(mylist)):
                if '(e.g., ' in mylist[i]:
                    mylist[i] = mylist[i].lstrip('(e.g., ').rstrip(')')
                elif 'e.g., ' in mylist[i]:
                    mylist[i] = mylist[i].lstrip('e.g., ')
            return mylist
        mylist = rmEG(mylist)
        
        mylist_index = list(range(len(mylist)))
        mylist_dict = dict(zip(mylist,mylist_index))
        
        # give citation its number
        def giveNum(A, B):
            for key in A.keys():
                B[key] = A[key]
            return B
        
        mylist_dict = giveNum(mydict, mylist_dict)
        
        # replacing in-text citation
        intx_refer = 0
        for intx_refer in range(len(para_l_c)):
            for key,value in mylist_dict.items():
                document.paragraphs[para_l_c[intx_refer]].text = document.paragraphs[para_l_c[intx_refer]].text.replace(key,('[' + str(value) + ']'))
            intx_refer = intx_refer + 1
        
        # redo font style
        total_para = para_l_c+para_l_c_r
        total_refer = 0
        for total_refer in range(len(total_para)):
            document.paragraphs[total_para[total_refer]].style.font.name = 'Times New Roman'
            total_refer = total_refer + 1
        
        newfilepath = filepath.rstrip('.docx')
        document.save(newfilepath+'_IEEE_style.docx')
        
    elif style == 'APA':
        # find index of references paragraph
        def findRefer():
            n = 0
            for para in document.paragraphs: 
                if para.text != 'References':
                    n = n + 1
                else:
                    return n
                
        refer_para = findRefer()
        # find out all subsets and return its index
        def find_all(source,dest):
            length1,length2 = len(source),len(dest)
            dest_list = []
            temp_list = []
            if length1 < length2:
                return -1
            i = 0
            while i <= length1-length2:
                if source[i] == dest[0]:
                    dest_list.append(i)
                i += 1
            if dest_list == []:
                return -1
            for x in dest_list:
                if source[x:x+length2] != dest:
                    temp_list.append(x)
            for x in temp_list:
                dest_list.remove(x)
            return dest_list
        
        # change reference list format
        citation_num = 0
        for each_refer in document.paragraphs[refer_para+1:]:
            if each_refer.text[-5:-3].isdigit():
                citation_num = citation_num + 1
                test_str = each_refer.text
                test_str = test_str[3:]
                test_str = test_str + ','
                split_test_str = test_str.split(',')
                def getAuthorName():
                    author_name_list = []
                    for i in split_test_str:
                        if '. ' in i:
                            if i.replace('. ', '').lstrip(' ').isalpha():
                                author_name_list.append(i.lstrip(' '))
                    return author_name_list
                author_name_list = getAuthorName()
                last_author_name = author_name_list[-1]
                last_au_name_index = test_str.find(last_author_name)
                year_index = last_au_name_index + len(last_author_name) + 1
                year = test_str[-6:-2]
                test_str = test_str[:year_index] + ' (' + year + '). ' + test_str[(year_index+1):]
                test_str = test_str[:-6]
                each_refer.text = test_str
                
        # transfering citation format
        def getCitaContent(refer_para):
            cita_total_list = []
            for each_refer in document.paragraphs[refer_para+1:]:
                if each_refer.text[:1].isalpha():
                    au_name_index = each_refer.text.find('(')
                    if each_refer.text[:au_name_index].count(', ') == 1:
                        cita_name = each_refer.text[:au_name_index]
                        cita_name = cita_name.rstrip(' ')
                        period_index = cita_name.rfind('.') + 2
                        cita_name = cita_name[period_index:]
                        cita_year_l = au_name_index + 1
                        cita_year_r = each_refer.text.find(')')
                        cita_year = each_refer.text[cita_year_l:cita_year_r]
                        cita_total = '(' + cita_name + ' ' + cita_year + ')'
                        cita_total_list.append(cita_total)
                    elif each_refer.text[:au_name_index].count(', ') == 2:
                        cita_name = each_refer.text[:au_name_index]
                        cita_name = cita_name.rstrip(' ')
                        period_index = find_all(cita_name, '.')
                        first_comma = cita_name.find(',')
                        cita_name_one = cita_name[period_index[0]+2:first_comma]
                        cita_name_two = cita_name[period_index[-1]+2:]
                        cita_year_l = au_name_index + 1
                        cita_year_r = each_refer.text.find(')')
                        cita_year = each_refer.text[cita_year_l:cita_year_r]
                        cita_total = '(' + cita_name_one + ' and ' + cita_name_two + ' ' + cita_year + ')'
                        cita_total_list.append(cita_total)
                    elif each_refer.text[:au_name_index].count(', ') >= 3:
                        first_comma = each_refer.text.find(',')
                        cita_name = each_refer.text[:first_comma]
                        period_index = find_all(cita_name, '.')
                        cita_name = cita_name[period_index[-1]+2:] + ' et al., '
                        cita_year_l = au_name_index + 1
                        cita_year_r = each_refer.text.find(')')
                        cita_year = each_refer.text[cita_year_l:cita_year_r]
                        cita_total = '(' + cita_name + cita_year + ')'
                        cita_total_list.append(cita_total)
            return cita_total_list
        
        cita_total_list = getCitaContent(refer_para)
        citation_num_list = list(range(1,citation_num+1))
        def citaNumChange():
            for i in range(len(citation_num_list)):
                citation_num_list[i] = '[' + str(citation_num_list[i]) + ']'
            return citation_num_list
        citation_num_list = citaNumChange()
        mylist_dict = dict(zip(citation_num_list, cita_total_list))
        # replacing in-text citation
        for intx_refer in range(refer_para):
            for key,value in mylist_dict.items():
                document.paragraphs[intx_refer].text = document.paragraphs[intx_refer].text.replace(key, value)
        # redo font style
        total_para = refer_para + citation_num + 1
        total_refer = 0
        for total_refer in range(total_para):
            document.paragraphs[total_refer].style.font.name = 'Times New Roman'
            total_refer = total_refer + 1
            
        newfilepath = filepath.rstrip('.docx')
        document.save(newfilepath+'_APA_style.docx')
        
        
change_style('C:/Z_STUDY/Master/1_Python for Data Analysis/Coursework/APAexample_no_hyperlinks.docx', 'IEEE')
change_style('C:/Z_STUDY/Master/1_Python for Data Analysis/Coursework/IEEEexample_no_hyperlinks.docx', 'APA')

        
'''        
# ---- DO NOT CHANGE THE CODE BELOW ----
if __name__ == "__main__":
    if len(sys.argv)<3: raise ValueError('Provide filename and style as input arguments')
    filepath, style = sys.argv[1], sys.argv[2]
    print('filepath is "{}"'.format(filepath))
    print('target style is "{}"'.format(style))
    change_style(filepath, style)
    '''